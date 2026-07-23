"""
从 Word 文档中提取文字和图片。
用法：python extract_word.py <docx路径> [输出图片目录]
输出：文字打印到 stdout，图片保存到指定目录（默认为 docx 同目录下的 sop_images/）
"""

import sys
import os

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')

try:
    import docx
except ImportError:
    print("错误：缺少 python-docx 依赖，请先执行：pip install python-docx", file=sys.stderr)
    sys.exit(1)

def extract(docx_path, img_dir=None):
    if not os.path.exists(docx_path):
        print(f"错误：文件不存在 {docx_path}", file=sys.stderr)
        sys.exit(1)

    if img_dir is None:
        img_dir = os.path.join(os.path.dirname(docx_path), 'sop_images')

    os.makedirs(img_dir, exist_ok=True)

    doc = docx.Document(docx_path)

    # 提取文字
    print("=== 文字内容 ===")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"{para.text}")

    # 提取表格
    if doc.tables:
        print("\n=== 表格内容 ===")
        for t_idx, table in enumerate(doc.tables):
            print(f"\n表格 {t_idx + 1}:")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                print(" | ".join(cells))

    # 提取图片
    img_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_count += 1
            img_data = rel.target_part.blob
            content_type = rel.target_part.content_type
            ext = content_type.split('/')[-1]
            if ext == 'jpeg':
                ext = 'jpg'
            img_path = os.path.join(img_dir, f'step_{img_count}.{ext}')
            with open(img_path, 'wb') as f:
                f.write(img_data)

    print(f"\n=== 图片提取 ===")
    print(f"共提取 {img_count} 张图片，保存到：{img_dir}")
    for i in range(1, img_count + 1):
        # 列出实际保存的文件
        for f in os.listdir(img_dir):
            if f.startswith(f'step_{i}.'):
                print(f"  图{i}：{os.path.join(img_dir, f)}")
                break


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法：python extract_word.py <docx路径> [输出图片目录]")
        sys.exit(1)

    docx_path = sys.argv[1]
    img_dir = sys.argv[2] if len(sys.argv) > 2 else None
    extract(docx_path, img_dir)
